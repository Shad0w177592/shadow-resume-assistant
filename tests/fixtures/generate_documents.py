from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "tests" / "fixtures" / "documents"


def make_text_pdf(path: Path, title: str, lines: list[str], pages: int = 1) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    for page in range(pages):
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(54, 790, f"{title} - {page + 1}")
        pdf.setFont("Helvetica", 11)
        y = 755
        for line in lines:
            pdf.drawString(54, y, line)
            y -= 20
        pdf.showPage()
    pdf.save()


def make_encrypted_pdf(source: Path, target: Path) -> None:
    reader = PdfReader(str(source))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("shadow-test")
    with target.open("wb") as stream:
        writer.write(stream)


def make_docx(path: Path, title: str, with_table: bool = False, with_list: bool = False) -> None:
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph("合成测试简历，不包含真实个人信息。")
    doc.add_heading("项目经历", 1)
    doc.add_paragraph("影子简历助手｜产品负责人｜2025.09-至今")
    if with_list:
        doc.add_paragraph("构建资料、岗位、生成和导出闭环。", style="List Bullet")
        doc.add_paragraph("使用固定测试集验证事实约束。", style="List Bullet")
    if with_table:
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "技能"
        table.cell(0, 1).text = "证据"
        table.cell(1, 0).text = "Python"
        table.cell(1, 1).text = "课程项目"
    doc.save(path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    make_text_pdf(OUT / "pdf-01-basic.pdf", "Resume", ["AI Agent Product", "OpenAI API", "2025.03-2025.08"])
    make_text_pdf(OUT / "pdf-02-multipage.pdf", "Resume", ["Project Experience", "Accuracy 71% to 88%"], pages=2)
    make_text_pdf(OUT / "pdf-03-table-like.pdf", "Resume", ["Skill | Evidence", "Python | Course project"])
    buffer = OUT / "pdf-encryption-source.pdf"
    make_text_pdf(buffer, "Encrypted Resume", ["Synthetic content only"])
    make_encrypted_pdf(buffer, OUT / "pdf-04-encrypted.pdf")
    buffer.unlink()
    (OUT / "pdf-05-damaged.pdf").write_bytes(b"%PDF-1.7\nnot-a-valid-pdf")

    make_docx(OUT / "docx-01-basic.docx", "合成简历一")
    make_docx(OUT / "docx-02-list.docx", "合成简历二", with_list=True)
    make_docx(OUT / "docx-03-table.docx", "合成简历三", with_table=True)
    make_docx(OUT / "docx-04-mixed.docx", "合成简历四", with_table=True, with_list=True)
    (OUT / "docx-05-damaged.docx").write_bytes(b"not-a-valid-docx")
    print(f"generated 10 document fixtures in {OUT}")


if __name__ == "__main__":
    main()

