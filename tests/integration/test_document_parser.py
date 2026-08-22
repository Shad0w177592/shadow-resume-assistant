from pathlib import Path

import pytest
from docx import Document

from app.domain.documents import ParseStatus
from app.services.document_parser import DocumentParser


ROOT = Path(__file__).resolve().parents[2]
FIXTURES = ROOT / "tests" / "fixtures" / "documents"


@pytest.fixture(scope="module", autouse=True)
def generate_fixtures() -> None:
    from tests.fixtures.generate_documents import main

    main()


@pytest.mark.integration
def test_all_ten_documents_parse_or_fail_with_classified_reason() -> None:
    parser = DocumentParser()
    files = sorted(FIXTURES.glob("*.pdf")) + sorted(FIXTURES.glob("*.docx"))
    assert len(files) == 10
    results = {path.name: parser.parse(path) for path in files}
    assert results["pdf-01-basic.pdf"].status == ParseStatus.PARSED
    assert results["pdf-02-multipage.pdf"].status == ParseStatus.PARSED
    assert len(results["pdf-02-multipage.pdf"].pages) == 2
    assert results["pdf-04-encrypted.pdf"].status == ParseStatus.ENCRYPTED
    assert results["pdf-05-damaged.pdf"].status == ParseStatus.DAMAGED
    assert results["docx-03-table.docx"].pages[0].blocks[-1].kind == "table"
    assert results["docx-05-damaged.docx"].status == ParseStatus.DAMAGED
    for result in results.values():
        assert result.status == ParseStatus.PARSED or result.failure_reason


@pytest.mark.integration
def test_txt_markdown_and_unsupported_formats_are_classified(tmp_path: Path) -> None:
    text = tmp_path / "resume.txt"
    markdown = tmp_path / "portfolio.md"
    unsupported = tmp_path / "resume.rtf"
    text.write_text("项目经历\n完成本地应用", encoding="utf-8")
    markdown.write_text("# 技能\nPython 与 React", encoding="utf-8")
    unsupported.write_text("data", encoding="utf-8")
    parser = DocumentParser()
    assert parser.parse(text).status == ParseStatus.PARSED
    parsed_markdown = parser.parse(markdown)
    assert parsed_markdown.pages[0].blocks[0].kind == "heading"
    assert parser.parse(unsupported).status == ParseStatus.UNSUPPORTED


@pytest.mark.integration
def test_docx_keeps_interleaved_table_heading_reading_order(tmp_path: Path) -> None:
    source = tmp_path / "interleaved.docx"
    document = Document()
    heading = document.add_table(rows=1, cols=2)
    heading.cell(0, 1).text = "在校经历"
    document.add_paragraph("2022.10-2023.10 再读轩书友会")
    second_heading = document.add_table(rows=1, cols=2)
    second_heading.cell(0, 1).text = "自我介绍"
    document.add_paragraph("愿意持续学习。")
    document.save(source)

    blocks = DocumentParser().parse(source).pages[0].blocks
    assert [block.kind for block in blocks] == ["table", "paragraph", "table", "paragraph"]
    assert [block.text.replace("|", "").strip() for block in blocks] == [
        "在校经历",
        "2022.10-2023.10 再读轩书友会",
        "自我介绍",
        "愿意持续学习。",
    ]
