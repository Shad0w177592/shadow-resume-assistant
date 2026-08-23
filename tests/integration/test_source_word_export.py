from __future__ import annotations

import hashlib
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.shared import Pt, RGBColor
from fastapi.testclient import TestClient

from app.main import create_app
from app.security.credentials import InMemoryCredentialStore
from app.services.source_word_export import SourceWordExport

HEADERS = {"x-shadow-session": "source-word-export"}


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _package_hashes(path: Path) -> dict[str, str]:
    with ZipFile(path) as archive:
        return {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name != "word/document.xml"
        }


def _experience(document: Document, company: str, body: str) -> None:
    date = document.add_paragraph()
    date.paragraph_format.space_after = Pt(3)
    date_run = date.add_run(f"2025.01-2025.06 {company}")
    date_run.bold = True
    date_run.font.color.rgb = RGBColor(31, 78, 120)
    role = document.add_paragraph()
    role_run = role.add_run("运营实习生")
    role_run.italic = True
    content = document.add_paragraph()
    content.paragraph_format.line_spacing = 1.25
    content_run = content.add_run(body)
    content_run.font.size = Pt(10.5)


def test_word_source_export_preserves_format_and_reorders_entries(
    monkeypatch, tmp_path: Path
) -> None:
    monkeypatch.setenv("SHADOW_SESSION_TOKEN", "source-word-export")
    source = tmp_path / "原排版简历.docx"
    original = Document()
    original.sections[0].header.paragraphs[0].text = "原简历页眉装饰"
    original.sections[0].footer.paragraphs[0].text = "原简历页脚"
    heading = original.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading_run = heading.add_run("工作经历")
    heading_run.bold = True
    heading_run.font.color.rgb = RGBColor(31, 78, 120)
    _experience(original, "甲公司", "负责甲公司的原始内容")
    _experience(original, "乙公司", "负责乙公司的原始内容")
    original.save(source)
    source_hash = _sha256(source)
    source_parts = _package_hashes(source)

    app = create_app(tmp_path / "data", InMemoryCredentialStore())
    with TestClient(app) as client:
        imported = client.post(
            "/api/imports/from-path", headers=HEADERS, json={"path": str(source)}
        ).json()
        assert [item["section_key"] for item in imported["candidates"]] == [
            "work",
            "work",
        ]
        confirmed = client.post(
            f"/api/imports/{imported['id']}/confirm",
            headers=HEADERS,
            json={
                "decisions": [
                    {"candidate_id": item["id"], "action": "accept"}
                    for item in imported["candidates"]
                ]
            },
        )
        assert confirmed.status_code == 200
        entries = client.get("/api/profile", headers=HEADERS).json()["entries"]
        for entry in entries:
            importance = 5 if "乙公司" in entry["title"] else 1
            updated = client.put(
                f"/api/profile/entries/{entry['id']}",
                headers=HEADERS,
                json={
                    "section_key": entry["section_key"],
                    "title": entry["title"],
                    "payload": entry["payload"],
                    "importance": importance,
                },
            )
            assert updated.status_code == 200

        job = client.post(
            "/api/jobs",
            headers=HEADERS,
            json={"jd_text": "乙公司相关岗位", "title": "运营", "company": "目标公司"},
        ).json()
        client.post(f"/api/jobs/{job['id']}/analyze", headers=HEADERS)
        config = client.get(
            f"/api/jobs/{job['id']}/resume-config", headers=HEADERS
        ).json()["config"]
        config["rewrite_sections"] = ["work"]
        saved = client.put(
            f"/api/jobs/{job['id']}/resume-config",
            headers=HEADERS,
            json={"config": config},
        )
        assert saved.status_code == 200, saved.text
        generated = client.post(f"/api/jobs/{job['id']}/generate", headers=HEADERS)
        assert generated.status_code == 200, generated.text
        draft = generated.json()
        blocks = draft["document"]["sections"][0]["blocks"]
        assert [block["heading"] for block in blocks] == ["乙公司", "甲公司"]
        for block in blocks:
            block["paragraphs"][0]["text"] = f"{block['heading']}针对岗位调整后的正文"
        saved = client.put(
            f"/api/jobs/{job['id']}/draft",
            headers=HEADERS,
            json={"document": draft["document"]},
        )
        assert saved.status_code == 200
        exported = client.post(
            f"/api/jobs/{job['id']}/export",
            headers=HEADERS,
            json={"filename": "保留原排版", "formats": ["docx"]},
        )
        assert exported.status_code == 200, exported.text
        assert exported.json()["word_mode"] == "source_format"
        target = Path(exported.json()["files"][0])

    assert _sha256(source) == source_hash
    assert _package_hashes(target) == source_parts
    result = Document(target)
    texts = [paragraph.text for paragraph in result.paragraphs]
    assert texts.index("2025.01-2025.06 乙公司") < texts.index("2025.01-2025.06 甲公司")
    assert "乙公司针对岗位调整后的正文" in texts
    assert "甲公司针对岗位调整后的正文" in texts
    exported_heading = next(
        paragraph for paragraph in result.paragraphs if paragraph.text == "工作经历"
    )
    assert exported_heading.runs[0].bold is True
    assert exported_heading.runs[0].font.color.rgb == RGBColor(31, 78, 120)
    assert result.sections[0].header.paragraphs[0].text == "原简历页眉装饰"
    assert result.sections[0].footer.paragraphs[0].text == "原简历页脚"


def test_overflow_is_reported_after_draft_is_saved(monkeypatch, tmp_path: Path) -> None:
    monkeypatch.setenv("SHADOW_SESSION_TOKEN", "source-word-export")
    app = create_app(tmp_path / "overflow-data", InMemoryCredentialStore())
    with TestClient(app) as client:
        entry = client.post(
            "/api/profile/entries",
            headers=HEADERS,
            json={
                "section_key": "project",
                "title": "长项目",
                "payload": {"content": "真实项目内容" * 400},
                "importance": 5,
            },
        ).json()
        job = client.post(
            "/api/jobs",
            headers=HEADERS,
            json={"jd_text": "项目经验", "title": "岗位", "company": "公司"},
        ).json()
        config = client.get(
            f"/api/jobs/{job['id']}/resume-config", headers=HEADERS
        ).json()["config"]
        config["entry_modes"][entry["id"]] = "must_include"
        saved = client.put(
            f"/api/jobs/{job['id']}/resume-config",
            headers=HEADERS,
            json={"config": config},
        )
        assert saved.status_code == 200
        generated = client.post(f"/api/jobs/{job['id']}/generate", headers=HEADERS)
        assert generated.status_code == 200, generated.text
        assert generated.json()["layout"]["status"] == "overflow"
        assert any("草稿已生成" in warning for warning in generated.json()["warnings"])
        persisted = client.get(f"/api/jobs/{job['id']}/draft", headers=HEADERS)
        assert persisted.status_code == 200


def test_source_word_export_only_changes_selected_summary_and_skills(
    monkeypatch, tmp_path: Path
) -> None:
    monkeypatch.setenv("SHADOW_SESSION_TOKEN", "source-word-export")
    source = tmp_path / "保留栏目简历.docx"
    original = Document()
    original.sections[0].header.paragraphs[0].text = "固定页眉"
    original.sections[0].footer.paragraphs[0].text = "固定页脚"
    for heading_text in ("工作经历", "校园经历", "专业技能", "自我介绍"):
        heading = original.add_table(rows=1, cols=2)
        heading.cell(0, 1).text = heading_text
        if heading_text == "工作经历":
            _experience(original, "原公司", "原工作内容不得被修改")
        elif heading_text == "校园经历":
            original.add_paragraph("2022.09-2023.06 辩论社")
            original.add_paragraph("队员")
            original.add_paragraph("原校园经历不得被修改")
        elif heading_text == "专业技能":
            original.add_paragraph("视频剪辑：熟练使用剪辑工具")
        else:
            original.add_paragraph("原始自我介绍")
    original.save(source)
    source_parts = _package_hashes(source)

    app = create_app(tmp_path / "selected-data", InMemoryCredentialStore())
    with TestClient(app) as client:
        imported = client.post(
            "/api/imports/from-path", headers=HEADERS, json={"path": str(source)}
        ).json()
        assert [item["section_key"] for item in imported["candidates"]] == [
            "work",
            "campus",
            "skills",
            "summary",
        ]
        confirmed = client.post(
            f"/api/imports/{imported['id']}/confirm",
            headers=HEADERS,
            json={
                "decisions": [
                    {"candidate_id": item["id"], "action": "accept"}
                    for item in imported["candidates"]
                ]
            },
        )
        assert confirmed.status_code == 200

        job = client.post(
            "/api/jobs",
            headers=HEADERS,
            json={
                "jd_text": "需要 AI 工具应用能力",
                "title": "AI 产品",
                "company": "目标公司",
            },
        ).json()
        client.post(f"/api/jobs/{job['id']}/analyze", headers=HEADERS)
        config = client.get(
            f"/api/jobs/{job['id']}/resume-config", headers=HEADERS
        ).json()["config"]
        config["rewrite_sections"] = ["summary", "skills"]
        saved = client.put(
            f"/api/jobs/{job['id']}/resume-config",
            headers=HEADERS,
            json={"config": config},
        )
        assert saved.status_code == 200
        generated = client.post(f"/api/jobs/{job['id']}/generate", headers=HEADERS)
        assert generated.status_code == 200, generated.text
        draft = generated.json()["document"]
        draft["sections"] = [
            section
            for section in draft["sections"]
            if section["section_key"] != "campus"
        ]
        work = next(
            section for section in draft["sections"] if section["section_key"] == "work"
        )
        assert work["blocks"][0]["heading"] == "原公司"
        assert any(
            paragraph["text"] == "原工作内容不得被修改"
            for paragraph in work["blocks"][0]["paragraphs"]
        )

        skills = next(
            section
            for section in draft["sections"]
            if section["section_key"] == "skills"
        )
        skills["blocks"].insert(
            0,
            {
                "block_id": "ai-skill",
                "heading": "AI 工具应用",
                "meta": "",
                "paragraphs": [
                    {
                        "paragraph_id": "ai-skill-p",
                        "text": "使用 Codex 完成项目从零到一搭建",
                        "source_entry_ids": [],
                        "risk_flags": ["ai_added_skill"],
                    }
                ],
            },
        )
        summary = next(
            section
            for section in draft["sections"]
            if section["section_key"] == "summary"
        )
        summary["blocks"][0]["paragraphs"][0]["text"] = "面向 AI 产品岗位的新版自我介绍"
        updated = client.put(
            f"/api/jobs/{job['id']}/draft",
            headers=HEADERS,
            json={"document": draft},
        )
        assert updated.status_code == 200, updated.text
        exported = client.post(
            f"/api/jobs/{job['id']}/export",
            headers=HEADERS,
            json={"filename": "仅修改勾选栏目", "formats": ["docx"]},
        )
        assert exported.status_code == 200, exported.text
        assert exported.json()["word_mode"] == "source_format"
        target = Path(exported.json()["files"][0])

        legacy_document = draft.copy()
        legacy_document["personal_info"] = dict(draft["personal_info"])
        legacy_document["personal_info"]["headline"] = "旧草稿字段中的自我介绍"
        legacy_document["sections"] = [
            section
            for section in draft["sections"]
            if section["section_key"] != "summary"
        ]
        legacy_saved = client.put(
            f"/api/jobs/{job['id']}/draft",
            headers=HEADERS,
            json={"document": legacy_document},
        )
        assert legacy_saved.status_code == 200, legacy_saved.text
        legacy_exported = client.post(
            f"/api/jobs/{job['id']}/export",
            headers=HEADERS,
            json={"filename": "旧草稿自我介绍兼容", "formats": ["docx"]},
        )
        assert legacy_exported.status_code == 200, legacy_exported.text
        assert legacy_exported.json()["word_mode"] == "source_format"
        legacy_target = Path(legacy_exported.json()["files"][0])

    assert _package_hashes(target) == source_parts
    result = Document(target)
    all_text = "\n".join(
        node.text or ""
        for node in result.element.body.iter()
        if node.tag.endswith("}t")
    )
    assert "原工作内容不得被修改" in all_text
    assert "原校园经历不得被修改" in all_text
    assert "面向 AI 产品岗位的新版自我介绍" in all_text
    assert "原始自我介绍" not in all_text
    assert all_text.index("AI 工具应用") < all_text.index("视频剪辑")
    assert result.sections[0].header.paragraphs[0].text == "固定页眉"
    assert result.sections[0].footer.paragraphs[0].text == "固定页脚"
    legacy_text = "\n".join(
        paragraph.text for paragraph in Document(legacy_target).paragraphs
    )
    assert "旧草稿字段中的自我介绍" in legacy_text
    assert "原始自我介绍" not in legacy_text


def test_duplicate_word_imports_merge_unique_source_blocks() -> None:
    entries = [
        {
            "id": "old-campus",
            "created_at": "1",
            "section_key": "campus",
            "source": {"document_id": "old", "block_ids": ["paragraph-8"]},
        },
        {
            "id": "old-work",
            "created_at": "2",
            "section_key": "work",
            "source": {"document_id": "old", "block_ids": ["paragraph-3"]},
        },
        {
            "id": "chosen-work",
            "created_at": "3",
            "section_key": "work",
            "source": {"document_id": "chosen", "block_ids": ["paragraph-3"]},
        },
        {
            "id": "chosen-skills",
            "created_at": "4",
            "section_key": "skills",
            "source": {"document_id": "chosen", "block_ids": ["paragraph-12"]},
        },
    ]

    merged = SourceWordExport._deduplicate_source_entries(entries, "chosen")

    assert [entry["id"] for entry in merged] == [
        "chosen-work",
        "chosen-skills",
        "old-campus",
    ]
