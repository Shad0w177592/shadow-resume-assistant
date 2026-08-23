from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import create_app
from app.security.credentials import InMemoryCredentialStore


HEADERS = {"x-shadow-session": "import-session"}


def test_import_candidates_require_confirmation_and_keep_source_locator(
    monkeypatch, tmp_path: Path
) -> None:
    monkeypatch.setenv("SHADOW_SESSION_TOKEN", "import-session")
    source = tmp_path / "我的作品集.md"
    source.write_text(
        "# 项目经历\n影子简历助手：完成本地简历生成流程\n个人爱好：摄影",
        encoding="utf-8",
    )
    app = create_app(tmp_path / "data", InMemoryCredentialStore())
    with TestClient(app) as client:
        imported = client.post(
            "/api/imports/from-path", headers=HEADERS, json={"path": str(source)}
        )
        assert imported.status_code == 201
        payload = imported.json()
        assert payload["status"] == "parsed"
        assert len(payload["candidates"]) == 2
        assert {item["confidence"] for item in payload["candidates"]} == {"clear"}
        before = client.get("/api/profile", headers=HEADERS).json()
        assert before["entries"] == []
        first, second = payload["candidates"]
        assert first["section_key"] == "project"
        assert first["title"] == "影子简历助手"
        confirmed = client.post(
            f"/api/imports/{payload['id']}/confirm",
            headers=HEADERS,
            json={
                "decisions": [
                    {"candidate_id": first["id"], "action": "accept"},
                    {"candidate_id": second["id"], "action": "ignore"},
                ]
            },
        )
        assert confirmed.json() == {"accepted": 1, "ignored": 1}
        entries = client.get("/api/profile", headers=HEADERS).json()["entries"]
        assert len(entries) == 1
        assert all(
            entry["payload"]["source"]["document_id"] == payload["id"]
            for entry in entries
        )


def test_duplicate_is_flagged_and_parser_failures_do_not_create_candidates(
    monkeypatch, tmp_path: Path
) -> None:
    monkeypatch.setenv("SHADOW_SESSION_TOKEN", "import-session")
    source = tmp_path / "resume.txt"
    source.write_text("项目经历：完成离线资料管理", encoding="utf-8")
    app = create_app(tmp_path / "data", InMemoryCredentialStore())
    with TestClient(app) as client:
        client.post(
            "/api/profile/entries",
            headers=HEADERS,
            json={
                "section_key": "project",
                "title": "项目",
                "payload": {"content": "项目经历：完成离线资料管理"},
            },
        )
        imported = client.post(
            "/api/imports/from-path", headers=HEADERS, json={"path": str(source)}
        ).json()
        assert imported["candidates"][0]["duplicate_of"] is not None

        damaged = tmp_path / "broken.pdf"
        damaged.write_bytes(b"not a pdf")
        failure = client.post(
            "/api/imports/from-path", headers=HEADERS, json={"path": str(damaged)}
        ).json()
        assert failure["status"] == "damaged"
        assert failure["candidates"] == []


def _add_section_heading(document: Document, title: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = ""
    table.cell(0, 1).text = title


def test_docx_headings_control_grouping_and_summary_destination(
    monkeypatch, tmp_path: Path
) -> None:
    monkeypatch.setenv("SHADOW_SESSION_TOKEN", "import-session")
    source = tmp_path / "章节化简历.docx"
    document = Document()
    _add_section_heading(document, "在校经历")
    document.add_paragraph("2022.10-2023.10 再读轩书友会")
    document.add_paragraph("新媒体部长")
    document.add_paragraph("1. 组织读书活动 10 场")
    document.add_paragraph("2023.10-2024.06 经济学院辩论队")
    document.add_paragraph("队员")
    document.add_paragraph("1. 参加校级比赛")
    _add_section_heading(document, "工作经历")
    document.add_paragraph("2024.07-2025.01 示例科技有限公司")
    document.add_paragraph("负责内容运营")
    _add_section_heading(document, "实习经历")
    document.add_paragraph("2025.02-2025.06 示例文化公司")
    document.add_paragraph("实习运营，完成直播项目")
    _add_section_heading(document, "自我介绍")
    document.add_paragraph("善于从数据中发现规律，也有实习和社团经历。")
    document.add_paragraph("愿意持续学习并快速补齐岗位所需技能。")
    document.save(source)

    app = create_app(tmp_path / "data", InMemoryCredentialStore())
    with TestClient(app) as client:
        result = client.post(
            "/api/imports/from-path", headers=HEADERS, json={"path": str(source)}
        ).json()
        candidates = result["candidates"]
        assert [item["section_key"] for item in candidates] == [
            "campus",
            "campus",
            "work",
            "internship",
            "summary",
        ]
        assert [item["title"] for item in candidates[:4]] == [
            "再读轩书友会",
            "经济学院辩论队",
            "示例科技有限公司",
            "示例文化公司",
        ]
        assert "再读轩书友会" in candidates[0]["payload"]["content"]
        assert "新媒体部长" in candidates[0]["payload"]["content"]
        assert "经济学院辩论队" in candidates[1]["payload"]["content"]
        assert "实习和社团经历" in candidates[-1]["payload"]["content"]
        assert candidates[-1]["title"] == "自我介绍"
        assert all(item["confidence"] == "clear" for item in candidates)

        response = client.post(
            f"/api/imports/{result['id']}/confirm",
            headers=HEADERS,
            json={
                "decisions": [
                    {"candidate_id": item["id"], "action": "accept"}
                    for item in candidates
                ]
            },
        )
        assert response.json() == {"accepted": 5, "ignored": 0}
        profile = client.get("/api/profile", headers=HEADERS).json()
        assert [entry["section_key"] for entry in profile["entries"]] == [
            "campus",
            "campus",
            "work",
            "internship",
        ]
        assert profile["personal_info"]["summary"] == (
            "善于从数据中发现规律，也有实习和社团经历。\n"
            "愿意持续学习并快速补齐岗位所需技能。"
        )
