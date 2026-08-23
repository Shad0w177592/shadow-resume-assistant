from pathlib import Path
import re
from uuid import uuid4
from zipfile import ZipFile

import pytest
from docx import Document
from pypdf import PdfReader

from app.domain.resume import (
    PersonalInfo,
    ResumeBlock,
    ResumeDocument,
    ResumeParagraph,
    ResumeSection,
    ResumeTemplate,
)
from app.services.resume_export import (
    _paginate_sections,
    _section_weight,
    to_docx,
    to_html,
    to_pdf,
)


def make_resume(template: ResumeTemplate, pages: int) -> ResumeDocument:
    sections = []
    for index, title in enumerate(["个人简介", "项目经历", "教育背景", "专业技能"]):
        sections.append(
            ResumeSection(
                section_id=f"section-{index}",
                section_key=f"section_{index}",
                title=title,
                order=index,
                column="left"
                if template == ResumeTemplate.TECHNICAL_DOUBLE_COLUMN and index > 1
                else "right",
                blocks=[
                    ResumeBlock(
                        block_id=f"block-{index}",
                        heading=f"{title}示例",
                        meta="2025.03-2025.08",
                        paragraphs=[
                            ResumeParagraph(
                                paragraph_id=f"paragraph-{index}",
                                text=f"使用证据完成{title}内容。",
                                source_entry_ids=[uuid4()],
                            )
                        ],
                    )
                ],
            )
        )
    return ResumeDocument(
        resume_id=uuid4(),
        template=template,
        page_target=pages,
        personal_info=PersonalInfo(
            name="李明",
            headline="AI Agent 产品方向",
            contacts=["liming@example.com", "杭州"],
        ),
        sections=sections,
    )


@pytest.mark.integration
@pytest.mark.parametrize("template", list(ResumeTemplate))
@pytest.mark.parametrize("pages", [1, 2])
def test_four_export_combinations_open_and_share_content(
    tmp_path: Path, template: ResumeTemplate, pages: int
) -> None:
    resume = make_resume(template, pages)
    docx_path = tmp_path / f"{template.value}-{pages}.docx"
    pdf_path = tmp_path / f"{template.value}-{pages}.pdf"
    html = to_html(resume)
    to_docx(resume, docx_path)
    to_pdf(resume, pdf_path)
    docx = Document(docx_path)
    pdf = PdfReader(str(pdf_path))
    docx_text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert len(pdf.pages) == pages
    for keyword in ["李明", "个人简介", "项目经历"]:
        assert keyword in html
        assert keyword in docx_text
        assert keyword in pdf_text

    assert "■" in docx_text
    assert abs(docx.sections[0].page_width.cm - 21) < 0.1
    assert abs(docx.sections[0].page_height.cm - 29.7) < 0.1
    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        assert "word/numbering.xml" in archive.namelist()
        if template == ResumeTemplate.TECHNICAL_DOUBLE_COLUMN:
            assert 'w:num="2"' in document_xml


def test_pdf_wraps_long_chinese_text_without_truncating_tail(tmp_path: Path) -> None:
    resume = make_resume(ResumeTemplate.SINGLE_COLUMN, 1)
    tail = "这是必须出现在导出文件中的结尾标记"
    resume.sections[0].blocks[0].paragraphs[0].text = (
        "完成需求分析和原型设计，" * 18 + tail
    )
    target = tmp_path / "long.pdf"
    to_pdf(resume, target)
    text = "\n".join(page.extract_text() or "" for page in PdfReader(target).pages)
    assert tail in re.sub(r"\s+", "", text)


def test_sparse_and_dense_word_exports_adapt_font_and_line_spacing(
    tmp_path: Path,
) -> None:
    sparse = make_resume(ResumeTemplate.SINGLE_COLUMN, 1)
    dense = make_resume(ResumeTemplate.SINGLE_COLUMN, 1)
    for section in dense.sections:
        section.blocks[0].paragraphs[0].text *= 90

    sparse_path = tmp_path / "sparse.docx"
    dense_path = tmp_path / "dense.docx"
    to_docx(sparse, sparse_path)
    to_docx(dense, dense_path)
    sparse_normal = Document(sparse_path).styles["Normal"]
    dense_normal = Document(dense_path).styles["Normal"]

    assert sparse_normal.font.size.pt > dense_normal.font.size.pt
    assert sparse_normal.paragraph_format.line_spacing > (
        dense_normal.paragraph_format.line_spacing
    )


def test_explicit_layout_polish_forces_spacious_export(tmp_path: Path) -> None:
    resume = make_resume(ResumeTemplate.SINGLE_COLUMN, 1)
    for section in resume.sections:
        section.blocks[0].paragraphs[0].text *= 90
    resume.layout_density = "expanded"
    target = tmp_path / "polished-layout.docx"
    to_docx(resume, target)
    normal = Document(target).styles["Normal"]
    assert normal.font.size.pt >= 10
    assert normal.paragraph_format.line_spacing >= 1.25


def test_two_page_partition_balances_content_weight() -> None:
    sections = []
    for index, text_length in enumerate([400, 320, 80, 80, 80, 80]):
        sections.append(
            ResumeSection(
                section_id=f"balanced-section-{index}",
                section_key=f"balanced_{index}",
                title=f"栏目{index}",
                order=index,
                blocks=[
                    ResumeBlock(
                        block_id=f"balanced-block-{index}",
                        heading=f"经历{index}",
                        paragraphs=[
                            ResumeParagraph(
                                paragraph_id=f"balanced-paragraph-{index}",
                                text="内容" * text_length,
                                source_entry_ids=[uuid4()],
                            )
                        ],
                    )
                ],
            )
        )

    pages = _paginate_sections(sections, 2)
    weights = [sum(_section_weight(section) for section in page) for page in pages]
    assert len(pages) == 2
    assert [section.order for page in pages for section in page] == list(range(6))
    assert weights[1] >= weights[0] * 0.5
    assert weights[0] >= weights[1] * 0.5


def test_two_page_export_rejects_an_almost_empty_second_page(tmp_path: Path) -> None:
    resume = make_resume(ResumeTemplate.SINGLE_COLUMN, 2)
    resume.sections = resume.sections[:1]
    with pytest.raises(ValueError, match="第二页内容不足"):
        to_docx(resume, tmp_path / "unbalanced.docx")
    with pytest.raises(ValueError, match="第二页内容不足"):
        to_pdf(resume, tmp_path / "unbalanced.pdf")
