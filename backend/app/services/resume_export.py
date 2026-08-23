from __future__ import annotations

from dataclasses import dataclass
from html import escape
from math import ceil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

from app.domain.resume import ResumeDocument, ResumeSection, ResumeTemplate

FONT = "Microsoft YaHei"
NAVY = RGBColor(0x17, 0x3B, 0x67)
MUTED = RGBColor(0x66, 0x70, 0x85)


@dataclass(frozen=True)
class LayoutMetrics:
    body_font: float
    line_spacing: float
    paragraph_after: float
    section_before: float
    pdf_body_font: float
    pdf_line_height: float
    pdf_block_gap: float


def _layout_metrics(resume: ResumeDocument) -> LayoutMetrics:
    characters_per_page = len(resume.plain_text()) / resume.page_target
    if resume.layout_density == "expanded" or characters_per_page < 900:
        return LayoutMetrics(10.2, 1.28, 4.0, 8.0, 9.8, 14.2, 6.0)
    if characters_per_page > 1800:
        return LayoutMetrics(8.8, 1.05, 1.5, 3.0, 8.6, 10.5, 2.5)
    return LayoutMetrics(9.4, 1.14, 2.5, 5.0, 9.2, 12.0, 4.0)


def _section_weight(section: ResumeSection) -> int:
    text_units = sum(
        max(1, ceil(len(paragraph.text) / 42))
        for block in section.blocks
        for paragraph in block.paragraphs
    )
    return 2 + len(section.blocks) * 2 + text_units


def _paginate_sections(
    sections: list[ResumeSection], page_target: int
) -> list[list[ResumeSection]]:
    if page_target == 1:
        return [sections]
    if len(sections) < 2:
        return [sections, []]
    weights = [_section_weight(section) for section in sections]
    total = sum(weights)
    split = min(
        range(1, len(sections)),
        key=lambda index: abs(sum(weights[:index]) - total / 2),
    )
    return [sections[:split], sections[split:]]


def _ensure_second_page_is_substantial(
    page_groups: list[list[ResumeSection]],
) -> None:
    if len(page_groups) != 2:
        return
    first_weight = sum(_section_weight(section) for section in page_groups[0])
    second_weight = sum(_section_weight(section) for section in page_groups[1])
    if second_weight * 2 < first_weight:
        raise ValueError("两页简历的第二页内容不足第一页的一半，请改为一页或先使用润色补充内容")


PDF_FONT = "STSong-Light"
if PDF_FONT not in pdfmetrics.getRegisteredFontNames():
    pdfmetrics.registerFont(UnicodeCIDFont(PDF_FONT))


def to_html(resume: ResumeDocument) -> str:
    section_html = []
    for section in sorted(resume.sections, key=lambda item: item.order):
        blocks = []
        for block in section.blocks:
            paragraphs = "".join(f"<li>{escape(item.text)}</li>" for item in block.paragraphs)
            blocks.append(
                f'<article data-block-id="{escape(block.block_id)}">'
                f"<h3>{escape(block.heading)}</h3>"
                f"<p>{escape(block.meta)}</p><ul>{paragraphs}</ul></article>"
            )
        section_html.append(
            f'<section data-column="{escape(section.column)}"><h2>{escape(section.title)}</h2>'
            f"{''.join(blocks)}</section>"
        )
    return (
        '<!doctype html><html lang="zh-CN"><meta charset="utf-8">'
        f'<body data-template="{resume.template.value}" data-pages="{resume.page_target}">'
        f"<header><h1>{escape(resume.personal_info.name)}</h1>"
        f"<p>{escape(resume.personal_info.headline)}</p>"
        f"<p>{escape(' | '.join(resume.personal_info.contacts))}</p></header>"
        f"{''.join(section_html)}</body></html>"
    )


def _set_font(run, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _set_columns(section, count: int) -> None:
    columns = section._sectPr.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        section._sectPr.append(columns)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), "340")


def _configure_styles(doc: Document, metrics: LayoutMetrics) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(metrics.body_font)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(metrics.paragraph_after)
    normal.paragraph_format.line_spacing = metrics.line_spacing
    for name, size, before, after, color in (
        ("Title", 24, 0, 2, NAVY),
        ("Heading 1", 11, metrics.section_before, 3, NAVY),
        ("Heading 2", 10, 2, 1, RGBColor(0x20, 0x27, 0x33)),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Resume Meta" not in doc.styles:
        meta = doc.styles.add_style("Resume Meta", WD_STYLE_TYPE.PARAGRAPH)
        meta.font.name = FONT
        meta.font.size = Pt(9)
        meta.font.color.rgb = MUTED
        meta._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        meta.paragraph_format.space_after = Pt(2)


def _add_header(doc: Document, resume: ResumeDocument, photo_path: Path | None = None) -> None:
    if photo_path:
        photo = doc.add_paragraph()
        photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        photo.paragraph_format.space_after = Pt(1)
        photo.add_run().add_picture(str(photo_path), height=Cm(1.55))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    _set_font(paragraph.add_run("■  "), 10, NAVY, True)
    _set_font(paragraph.add_run(resume.personal_info.name or "姓名"), 24, NAVY, True)
    _set_font(paragraph.add_run("  ■"), 10, NAVY, True)
    if resume.personal_info.headline:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(2)
        _set_font(sub.add_run(resume.personal_info.headline), 9.5, MUTED)
    contacts = doc.add_paragraph()
    contacts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contacts.paragraph_format.space_after = Pt(4)
    _set_font(contacts.add_run("  |  ".join(resume.personal_info.contacts)), 9, MUTED)


def _add_resume_section(
    doc: Document, resume_section: ResumeSection, metrics: LayoutMetrics
) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    _set_font(heading.add_run("■  "), 6, NAVY, True)
    _set_font(heading.add_run(resume_section.title), 11, NAVY, True)
    borders = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "6")
    border.set(qn("w:color"), "DCE8F4")
    borders.append(border)
    heading._p.get_or_add_pPr().append(borders)
    for block in resume_section.blocks:
        title = doc.add_paragraph(style="Heading 2")
        title.paragraph_format.keep_with_next = True
        title.paragraph_format.tab_stops.add_tab_stop(Cm(18.4), WD_TAB_ALIGNMENT.RIGHT)
        _set_font(title.add_run(block.heading), 10, RGBColor(0x20, 0x27, 0x33), True)
        if block.meta:
            _set_font(title.add_run(f"\t{block.meta}"), 9, MUTED)
        for item in block.paragraphs:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(metrics.paragraph_after)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = metrics.line_spacing
            paragraph.paragraph_format.widow_control = True
            _set_font(paragraph.add_run(item.text), metrics.body_font, RGBColor(0x20, 0x27, 0x33))


def to_docx(resume: ResumeDocument, target: Path, *, photo_path: Path | None = None) -> None:
    doc = Document()
    metrics = _layout_metrics(resume)
    base = doc.sections[0]
    base.page_width = Cm(21)
    base.page_height = Cm(29.7)
    base.top_margin = Cm(1.2)
    base.bottom_margin = Cm(1.2)
    base.left_margin = Cm(1.3)
    base.right_margin = Cm(1.3)
    base.header_distance = Cm(0.6)
    base.footer_distance = Cm(0.6)
    _configure_styles(doc, metrics)
    _add_header(doc, resume, photo_path)
    content_section = base
    if resume.template == ResumeTemplate.TECHNICAL_DOUBLE_COLUMN:
        content_section = doc.add_section(WD_SECTION.CONTINUOUS)
        content_section.page_width = Cm(21)
        content_section.page_height = Cm(29.7)
        content_section.top_margin = Cm(1.2)
        content_section.bottom_margin = Cm(1.2)
        content_section.left_margin = Cm(1.3)
        content_section.right_margin = Cm(1.3)
        _set_columns(content_section, 2)
    sections = sorted(resume.sections, key=lambda item: item.order)
    page_groups = _paginate_sections(sections, resume.page_target)
    _ensure_second_page_is_substantial(page_groups)
    split = len(page_groups[0])
    for index, resume_section in enumerate(sections):
        if resume.page_target == 2 and index == split:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif (
            resume.template == ResumeTemplate.TECHNICAL_DOUBLE_COLUMN
            and index > 0
            and resume_section.column != sections[index - 1].column
        ):
            doc.add_paragraph().add_run().add_break(WD_BREAK.COLUMN)
        _add_resume_section(doc, resume_section, metrics)
    if resume.page_target == 2:
        header = content_section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_font(header.add_run(f"{resume.personal_info.name} | 简历"), 8.5, MUTED)
        footer = content_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_font(footer.add_run("2 / 2"), 8.5, MUTED)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


def _wrap(text: str, font: str, size: float, width: float) -> list[str]:
    lines = []
    current = ""
    for character in text:
        candidate = current + character
        if current and (
            len(current) >= 38 or pdfmetrics.stringWidth(candidate, font, size) > width
        ):
            lines.append(current)
            current = character
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines or [""]


def _draw_pdf_section(
    pdf, section: ResumeSection, x: float, y: float, width: float, metrics: LayoutMetrics
) -> float:
    pdf.setFillColor(HexColor("#173B67"))
    pdf.setFont(PDF_FONT, 11)
    pdf.drawString(x, y, f"■  {section.title}")
    pdf.setStrokeColor(HexColor("#DCE8F4"))
    pdf.line(x, y - 4, x + width, y - 4)
    y -= 18
    for block in section.blocks:
        pdf.setFillColor(HexColor("#202733"))
        pdf.setFont(PDF_FONT, 9.8)
        pdf.drawString(x, y, block.heading)
        if block.meta:
            pdf.setFillColor(HexColor("#667085"))
            pdf.setFont(PDF_FONT, 8.7)
            meta_width = pdfmetrics.stringWidth(block.meta, PDF_FONT, 8.7)
            pdf.drawString(x + width - meta_width, y, block.meta)
        y -= 14
        for paragraph in block.paragraphs:
            lines = _wrap(paragraph.text, PDF_FONT, metrics.pdf_body_font, width - 12)
            for line_index, line in enumerate(lines):
                pdf.setFillColor(HexColor("#202733"))
                pdf.setFont(PDF_FONT, metrics.pdf_body_font)
                pdf.drawString(x + 10, y, ("• " if line_index == 0 else "  ") + line)
                y -= metrics.pdf_line_height
        y -= metrics.pdf_block_gap
    return y - 5


def to_pdf(resume: ResumeDocument, target: Path, *, photo_path: Path | None = None) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(target), pagesize=A4)
    width, height = A4
    metrics = _layout_metrics(resume)
    sections = sorted(resume.sections, key=lambda item: item.order)
    page_groups = _paginate_sections(sections, resume.page_target)
    _ensure_second_page_is_substantial(page_groups)
    for page_index, page_sections in enumerate(page_groups):
        if page_index == 0:
            if photo_path:
                pdf.drawImage(
                    str(photo_path),
                    37,
                    height - 69,
                    width=30,
                    height=38,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            pdf.setFillColor(HexColor("#173B67"))
            pdf.setFont(PDF_FONT, 10)
            name_width = pdfmetrics.stringWidth(resume.personal_info.name, PDF_FONT, 22)
            center = width / 2
            pdf.drawString(center - name_width / 2 - 18, height - 40, "■")
            pdf.setFont(PDF_FONT, 22)
            pdf.drawString(center - name_width / 2, height - 40, resume.personal_info.name)
            pdf.setFont(PDF_FONT, 10)
            pdf.drawString(center + name_width / 2 + 8, height - 40, "■")
            pdf.setFillColor(HexColor("#667085"))
            pdf.setFont(PDF_FONT, 9)
            contact = "  |  ".join(resume.personal_info.contacts)
            contact_width = pdfmetrics.stringWidth(contact, PDF_FONT, 9)
            pdf.drawString(center - contact_width / 2, height - 58, contact)
            top = height - 84
        else:
            pdf.setFillColor(HexColor("#667085"))
            pdf.setFont(PDF_FONT, 8.5)
            pdf.drawRightString(width - 37, height - 25, f"{resume.personal_info.name} | 简历")
            top = height - 48
        if resume.template == ResumeTemplate.SINGLE_COLUMN:
            y = top
            for resume_section in page_sections:
                y = _draw_pdf_section(pdf, resume_section, 37, y, width - 74, metrics)
                if y < 35:
                    raise RuntimeError("PDF content overflow")
        else:
            left_x, gap, left_width = 37, 17, 150
            right_x = left_x + left_width + gap
            right_width = width - 37 - right_x
            left_y = right_y = top
            for resume_section in page_sections:
                if resume_section.column == "left":
                    left_y = _draw_pdf_section(
                        pdf, resume_section, left_x, left_y, left_width, metrics
                    )
                else:
                    right_y = _draw_pdf_section(
                        pdf, resume_section, right_x, right_y, right_width, metrics
                    )
            if min(left_y, right_y) < 35:
                raise RuntimeError("PDF content overflow")
        pdf.setFillColor(HexColor("#667085"))
        pdf.setFont(PDF_FONT, 8.5)
        if resume.page_target == 2:
            pdf.drawRightString(width - 37, 20, f"{page_index + 1} / 2")
        pdf.showPage()
    pdf.save()
    if len(PdfReader(str(target)).pages) != resume.page_target:
        raise RuntimeError("PDF page count validation failed")
