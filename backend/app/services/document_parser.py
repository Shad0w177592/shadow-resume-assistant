from __future__ import annotations

import mimetypes
from pathlib import Path
from uuid import uuid4
from zipfile import BadZipFile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.domain.documents import ParsedBlock, ParsedDocument, ParsedPage, ParseStatus

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DocumentParser:
    def parse(self, source: Path) -> ParsedDocument:
        suffix = source.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(source)
        if suffix == ".docx":
            return self._parse_docx(source)
        if suffix in {".txt", ".md"}:
            return self._parse_text(source, suffix)
        return self._failure(source, ParseStatus.UNSUPPORTED, "unsupported file type")

    def _parse_text(self, source: Path, suffix: str) -> ParsedDocument:
        try:
            text = source.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError:
            try:
                text = source.read_text(encoding="gb18030")
            except (UnicodeDecodeError, OSError) as exc:
                return self._failure(
                    source, ParseStatus.DAMAGED, f"unreadable text: {type(exc).__name__}"
                )
        except OSError as exc:
            return self._failure(
                source, ParseStatus.DAMAGED, f"unreadable text: {type(exc).__name__}"
            )
        blocks = []
        for order, line in enumerate(line.strip() for line in text.splitlines() if line.strip()):
            kind = "heading" if suffix == ".md" and line.startswith("#") else "paragraph"
            blocks.append(
                ParsedBlock(
                    block_id=f"line-{order}",
                    kind=kind,
                    text=line.lstrip("# ") if kind == "heading" else line,
                    order=order,
                )
            )
        return ParsedDocument(
            document_id=uuid4(),
            source_name=source.name,
            media_type="text/markdown" if suffix == ".md" else "text/plain",
            pages=[ParsedPage(page_number=1, blocks=blocks)],
            status=ParseStatus.PARSED,
        )

    def _parse_pdf(self, source: Path) -> ParsedDocument:
        try:
            reader = PdfReader(str(source))
            if reader.is_encrypted:
                return self._failure(source, ParseStatus.ENCRYPTED, "PDF is encrypted")
            pages: list[ParsedPage] = []
            has_text = False
            for page_number, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                has_text = has_text or bool(text)
                blocks = (
                    [
                        ParsedBlock(
                            block_id=f"page-{page_number}-paragraph-0",
                            kind="paragraph",
                            text=text,
                            order=0,
                        )
                    ]
                    if text
                    else []
                )
                pages.append(ParsedPage(page_number=page_number, blocks=blocks))
            status = ParseStatus.PARSED if has_text else ParseStatus.SCANNED
            reason = None if has_text else "PDF contains no selectable text"
            return ParsedDocument(
                document_id=uuid4(),
                source_name=source.name,
                media_type="application/pdf",
                pages=pages,
                status=status,
                failure_reason=reason,
            )
        except (PdfReadError, OSError, ValueError) as exc:
            return self._failure(source, ParseStatus.DAMAGED, f"damaged PDF: {type(exc).__name__}")

    def _parse_docx(self, source: Path) -> ParsedDocument:
        try:
            doc = Document(source)
            blocks: list[ParsedBlock] = []
            order = 0
            paragraph_index = 0
            table_index = 0
            # doc.paragraphs followed by doc.tables destroys the visual reading
            # order. Resume section headings are often placed in small tables,
            # so retain the real OOXML body order.
            for item in doc.iter_inner_content():
                if isinstance(item, Paragraph):
                    text = item.text.strip()
                    block_id = f"paragraph-{paragraph_index}"
                    paragraph_index += 1
                    if not text:
                        continue
                    kind = "list_item" if item.style and "List" in item.style.name else "paragraph"
                    if item.style and item.style.name.startswith("Heading"):
                        kind = "heading"
                    blocks.append(ParsedBlock(block_id=block_id, kind=kind, text=text, order=order))
                elif isinstance(item, Table):
                    rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
                    text = "\n".join(" | ".join(row) for row in rows).strip()
                    block_id = f"table-{table_index}"
                    table_index += 1
                    if not text:
                        continue
                    blocks.append(
                        ParsedBlock(
                            block_id=block_id,
                            kind="table",
                            text=text,
                            order=order,
                            table_rows=rows,
                        )
                    )
                else:
                    continue
                order += 1
            return ParsedDocument(
                document_id=uuid4(),
                source_name=source.name,
                media_type=DOCX_MIME,
                pages=[ParsedPage(page_number=1, blocks=blocks)],
                status=ParseStatus.PARSED,
            )
        except (BadZipFile, KeyError, OSError, ValueError) as exc:
            return self._failure(source, ParseStatus.DAMAGED, f"damaged DOCX: {type(exc).__name__}")

    @staticmethod
    def _failure(source: Path, status: ParseStatus, reason: str) -> ParsedDocument:
        media_type = mimetypes.guess_type(source.name)[0] or "application/octet-stream"
        return ParsedDocument(
            document_id=uuid4(),
            source_name=source.name,
            media_type=media_type,
            pages=[],
            status=status,
            failure_reason=reason,
        )
